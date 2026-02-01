"""Script to extract all content from the Values and Principles.docx file"""
from docx import Document

def extract_full_content(docx_path):
    doc = Document(docx_path)
    
    print("=" * 80)
    print("PARAGRAPHS")
    print("=" * 80)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text}")
    
    print("\n" + "=" * 80)
    print("TABLES")
    print("=" * 80)
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {t_idx + 1} ---")
        for row_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"Row {row_idx}: {row_text}")

if __name__ == "__main__":
    extract_full_content("Values and Principles.docx")
